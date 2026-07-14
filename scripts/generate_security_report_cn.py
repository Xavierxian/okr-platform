from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

from generate_security_report import (
    BLUE, DARK_BLUE, LIGHT_GRAY, MUTED, RED,
    add_bullet, add_heading, add_number, add_para, add_table, add_header_footer,
    configure_styles, set_cell_shading, set_run_font, set_table_geometry, style_cell,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "OKR平台安全修复报告_测试复测版_20260714.docx"


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(section)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("OKR 平台安全修复报告")
    set_run_font(run, size=24, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("安全初测问题整改情况 - 提交安全复测")
    set_run_font(run, size=12, color=MUTED)

    metadata = [
        ("系统名称", "OKR 平台"),
        ("报告日期", "2026-07-14"),
        ("整改范围", "安全初测报告八类问题及源码核验发现项"),
        ("提测目的", "验证服务端授权、认证安全基线、审计及 HTTPS 部署控制"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for label, value in metadata:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        style_cell(row.cells[0], label, bold=True, color=DARK_BLUE, size=10)
        style_cell(row.cells[1], value, size=10)

    add_heading(doc, "一、整改结论", 1)
    add_para(doc, "本次整改已在服务端建立“角色 + 资源关系”的统一授权机制，重点封堵了通过修改 URL、请求体和资源 ID 越权访问或修改其他用户 OKR 的风险。同时完成了输入校验、管理员认证收敛、CSRF 与来源校验、安全响应头、审计记录以及直接 Node.js HTTPS 启动能力的代码实现。")
    add_para(doc, "结论：高风险问题已完成代码级修复。生产证书部署、外网 TLS 扫描、钉钉认证联调和全角色回归测试尚需在测试环境完成后，方可判定生产发布通过。", bold_prefix="结论：", color=DARK_BLUE)

    add_heading(doc, "二、整改范围与已核验材料", 1)
    for text in [
        "源码核验范围：Express 路由、授权服务、认证与 Session、数据库结构及迁移、前端请求封装、应用启动与部署配置。",
        "自动化检查：TypeScript 类型检查、安全策略单元测试、代码规范检查；安全策略测试共 3 项，均已通过。",
        "开发环境运行核验：数据库迁移成功执行，应用进程重启后健康检查、首页及前端主资源均返回 HTTP 200。",
        "未在本机完成的验证：生产证书文件、公开域名、PM2 证书重载、外网 TLS 扫描，以及各业务角色测试账号的全量回归。",
    ]:
        add_bullet(doc, text)

    doc.add_page_break()
    add_heading(doc, "三、安全问题修复对照表", 1)
    rows = [
        ("1", "明文 HTTP 传输", "已完成代码实现，待部署验证", "生产环境由 Node.js HTTPS Server 承载，最低 TLS 1.2；独立 HTTP 端口统一 308 跳转至固定公开 HTTPS 地址。证书、私钥或公开地址无效时启动失败，不回退 HTTP。"),
        ("2", "CSRF", "已修复", "新增会话绑定 CSRF Token；所有状态变更接口必须携带 X-CSRF-Token；非白名单 Origin 被拒绝。"),
        ("3", "表单与参数篡改", "已修复", "使用 Zod DTO、字段白名单和范围校验；服务端以当前登录用户重新判定权限，不信任客户端传入的身份、部门或资源归属。"),
        ("4", "管理员弱口令", "已修复", "最少 8 位且包含四类字符中的至少三类；bcrypt 12 轮哈希；删除固定默认密码。"),
        ("5", "缺少安全响应头", "已修复", "配置 Helmet、CSP、防嵌入、X-Content-Type-Options、Referrer-Policy、生产 HSTS；敏感响应禁止缓存。"),
        ("6", "水平越权", "已修复", "目标、KR、评论、通知、导入导出、统计与 AI 分析均接入统一授权。无权资源返回 404，防止资源枚举。"),
        ("7", "垂直越权", "已修复", "管理员接口保留服务端超级管理员校验；前端管理路由增加角色守卫，普通用户跳转首页。"),
        ("8", "审计不足", "已修复", "新增 audit_logs，记录操作者、角色、时间、IP、UA、操作、资源、结果及请求 ID；仅超级管理员可查询和导出。"),
    ]
    add_table(doc, ["序号", "初测问题", "整改状态", "修复说明"], rows, [520, 1750, 2200, 4890])

    add_heading(doc, "四、核心安全控制说明", 1)
    add_heading(doc, "4.1 资源授权", 2)
    add_para(doc, "权限由服务端统一判定，不再由前端隐藏按钮或客户端字段决定。权限矩阵如下：")
    add_table(doc, ["身份", "允许的资源操作"], [
        ("超级管理员", "管理全部资源与系统配置。"),
        ("目标创建人", "管理本人创建的目标及其 KR。"),
        ("KR 执行人", "查看所属目标；仅更新本人 KR 进度和自评。"),
        ("KR 协同人", "查看所属目标和 KR；可发表评论。"),
        ("同部门成员 / 负责人", "按既有可见范围只读并发表评论；不得管理非本人创建资源。"),
        ("非管理员用户", "仅允许钉钉认证；账号密码登录会被拒绝。"),
    ], [2500, 6960])
    add_para(doc, "对当前用户无权访问的目标、KR、评论及通知资源，接口返回 404；访问管理员模块但角色不满足时返回 403。", color=DARK_BLUE)

    add_heading(doc, "4.2 认证、会话与密码", 2)
    for text in [
        "POST /api/auth/login 仅允许本地超级管理员登录；非管理员账号密码登录统一返回通用凭据错误信息。",
        "管理员密码存储为 bcrypt 哈希。旧版 AES 解密和启动时同步管理员密码的逻辑已移除。",
        "生产环境必须提供 SESSION_SECRET，不再存在固定后备值；本地登录成功后重新生成 Session ID。",
        "登录限制同时按 IP 与账号维度执行；连续失败触发短时锁定。",
        "生产 Cookie 强制 Secure、HttpOnly、SameSite=Lax；直接 TLS 部署不默认信任代理。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4.3 输入、浏览器与响应保护", 2)
    for text in [
        "状态变更接口使用 Zod DTO 与字段白名单，校验 ID、日期、角色、状态、0-100 进度、评分、权重、数组长度及导入行数。",
        "前端通过 GET /api/auth/csrf-token 获取 Token，并在状态变更请求中发送 X-CSRF-Token。",
        "Helmet 提供 CSP、防嵌入、X-Content-Type-Options、Referrer-Policy 与生产 HSTS；敏感页面响应设置 Cache-Control: no-store。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "4.4 审计与日志", 2)
    add_para(doc, "审计范围包括登录成功/失败、退出、改密、用户角色变更、部门和周期变更、目标/KR 操作、进度、评分、导入导出、AI 分析和清空数据。审计日志不记录密码、密码哈希、Session、CSRF Token、Cookie、完整请求体、证书路径和私钥信息。")
    add_para(doc, "普通请求日志改为结构化记录，仅包含请求 ID、方法、路径、状态码和耗时。审计记录不提供修改接口，保留策略为 180 天。")

    doc.add_page_break()
    add_heading(doc, "五、HTTPS 部署验收要求", 1)
    add_para(doc, "以下配置属于生产环境上线前置条件，测试环境需按相同模式验证：")
    add_table(doc, ["配置项", "验收要求"], [
        ("HTTPS_CERT_PATH", "指向可读取的 PEM 完整证书链。"),
        ("HTTPS_KEY_PATH", "指向可读取的 PEM 私钥；权限仅限运行账户读取。"),
        ("HTTPS_PORT", "HTTPS 监听端口，默认 5000。"),
        ("HTTP_PORT", "HTTP 跳转端口，默认 5001。"),
        ("PUBLIC_HTTPS_ORIGIN", "固定的公开 HTTPS 地址；跳转目标不得根据 Host Header 拼接。"),
        ("启动失败策略", "生产环境证书、私钥、公开地址缺失、不可读、格式错误或不匹配时必须退出。"),
        ("证书续期", "运维续期后执行 pm2 reload okr-platform 加载新证书。"),
    ], [2500, 6860])

    add_heading(doc, "六、安全复测清单", 1)
    add_heading(doc, "6.1 授权与越权", 2)
    for text in [
        "准备超级管理员、目标创建人、KR 执行人、协同人、同部门只读成员和跨部门成员账号。",
        "篡改 objectiveId、KR ID、评论 ID、通知 ID 和批量排序请求 ID，确认跨范围查询与操作均被拒绝。",
        "确认 KR 执行人不能修改他人 KR、目标创建人、部门、协同范围，或管理无关目标的排序。",
        "确认通知只能由接收人标记已读；评论仅作者或超级管理员可删除。",
        "为 AI 分析或排名传入任意部门参数，确认结果仅使用当前用户可见目标。",
    ]:
        add_number(doc, text)
    add_heading(doc, "6.2 认证与浏览器安全", 2)
    for text in [
        "以非管理员身份调用账号密码登录，确认得到通用凭据错误且不能创建会话。",
        "尝试弱管理员密码，确认被策略拒绝；对账号和 IP 连续失败，确认触发临时锁定。",
        "确认登录成功后 Session ID 发生变化；生产 HTTPS 下 Cookie 包含 Secure、HttpOnly、SameSite=Lax。",
        "对状态变更接口分别不带 CSRF Token、带错误 Token、使用非白名单 Origin 发起请求，确认均被拒绝。",
        "确认生产响应头包含 CSP、防嵌入、X-Content-Type-Options、Referrer-Policy 和 HSTS。",
    ]:
        add_number(doc, text)
    add_heading(doc, "6.3 HTTPS、审计与运维", 2)
    for text in [
        "使用有效 PEM 证书与私钥启动生产模式，确认 HTTPS 登录和 API 正常。",
        "分别使用缺失、不可读、非法或证书私钥不匹配的 TLS 文件启动，确认应用拒绝启动。",
        "验证 TLS 1.0/1.1 被拒绝；HTTP 返回 308，且路径和查询参数保持不变。",
        "注入伪造 Host Header，确认重定向目标不受影响；执行 PM2 reload 后确认新证书生效。",
        "检查登录、改密、用户变更、目标/KR、导入导出、AI 分析与清空数据均生成审计记录，且审计内容不含敏感值。",
    ]:
        add_number(doc, text)

    add_heading(doc, "七、上线条件与遗留事项", 1)
    add_table(doc, ["上线条件", "责任方", "验收材料"], [
        ("测试环境完成多角色安全回归", "测试", "复测用例执行记录、接口请求/响应证据、缺陷关闭记录。"),
        ("部署证书及生产环境变量", "运维", "PM2 配置、证书权限、HTTPS 访问验证材料。"),
        ("轮换已暴露凭据并清理旧 Secret", "运维 / 管理员", "新密钥托管记录、旧 AES 启动配置删除、历史 Session 失效证明。"),
        ("验证审计清理和监控", "后端 / 运维", "180 天清理任务记录及 TLS、登录失败、越权、异常、审计写入失败告警。"),
        ("处理三方依赖扫描项", "后端", "依赖升级或缓解方案；对暂无上游修复项记录残余风险。"),
    ], [3300, 1600, 4460])
    add_para(doc, "在上述上线条件完成、测试团队出具安全复测通过结论前，不得将本版本发布至生产环境。", bold_prefix="在上述上线条件完成", color=RED)

    add_heading(doc, "八、接口变更说明", 1)
    add_table(doc, ["接口或行为", "测试预期"], [
        ("GET /api/auth/csrf-token", "返回当前 Session 对应的 CSRF Token。"),
        ("POST /api/auth/login", "仅本地超级管理员可使用；无效凭据返回通用失败信息。"),
        ("GET /api/auth/me", "返回 authProvider，不返回密码相关字段。"),
        ("状态变更接口", "必须携带 X-CSRF-Token 且 Origin 在白名单内；未知字段被拒绝。"),
        ("资源访问接口", "无权资源返回 404；管理员模块无权限访问返回 403。"),
    ], [3000, 6360])

    doc.core_properties.title = "OKR 平台安全修复报告"
    doc.core_properties.subject = "安全复测材料"
    doc.core_properties.author = "OKR 平台研发"
    doc.core_properties.comments = "用于安全复测。"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

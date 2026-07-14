from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "安全修复报告_测试复测版_20260714.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
RED = "9B1C1C"
GREEN = "1F3A5F"
MUTED = "595959"


def set_run_font(run, size=11, bold=None, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell(cell, text, bold=False, color="000000", size=9.2, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        style_cell(cell, text, bold=True, color=DARK_BLUE, center=True)
    for index, row_data in enumerate(rows):
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            if index % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            style_cell(cell, text, center=(len(text) <= 8))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_para(doc, text="", bold_prefix=None, color="000000", after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        prefix = p.add_run(bold_prefix)
        set_run_font(prefix, bold=True, color=color)
        remainder = p.add_run(text[len(bold_prefix):])
        set_run_font(remainder, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, before, after, color in ((1, 16, 16, 8, BLUE), (2, 13, 12, 6, BLUE), (3, 12, 8, 4, DARK_BLUE)):
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("OKR Platform | Security Remediation Report")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Internal use only | Submitted for security regression testing | 2026-07-14")
    set_run_font(r, size=8.5, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_header_footer(section)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("OKR Platform Security Remediation Report")
    set_run_font(run, size=24, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Security assessment findings remediation status | Submitted for regression testing")
    set_run_font(run, size=12, color=MUTED)

    metadata = [
        ("System", "OKR Platform"),
        ("Report date", "2026-07-14"),
        ("Remediation scope", "Initial security test report and source-code verification findings"),
        ("Testing request", "Validate server-side authorization, authentication baseline, audit, and HTTPS deployment controls"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    for label, value in metadata:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        style_cell(row.cells[0], label, bold=True, color=DARK_BLUE, size=10)
        style_cell(row.cells[1], value, size=10)

    add_heading(doc, "1. Executive Conclusion", 1)
    add_para(doc, "This remediation closes the primary horizontal privilege-escalation path by enforcing authorization on the server using role and resource relationships. The application now validates write payloads, limits local password login to super administrators, introduces CSRF and origin controls, adds baseline HTTP security headers, and records security-relevant actions in an immutable audit log.")
    add_para(doc, "Assessment result: code-level remediation is complete for the identified high-risk areas. Production certificate deployment, external TLS scanning, DingTalk integration verification, and full multi-role regression remain release-gate test items and must be validated in the test environment before production release.", bold_prefix="Assessment result: ", color=DARK_BLUE)

    add_heading(doc, "2. Scope and Evidence", 1)
    add_bullet(doc, "Source review: Express routes, authorization helpers, authentication/session handling, database schema and migrations, client request handling, and operational startup configuration.")
    add_bullet(doc, "Automated checks performed during remediation: TypeScript type check, Node security-policy tests, and code linting. The security-policy test suite passed 3 of 3 tests.")
    add_bullet(doc, "Runtime verification performed in the development environment: database migration applied successfully; process restarted; health endpoint, root page, and main JavaScript asset returned HTTP 200.")
    add_bullet(doc, "Excluded from completed evidence: production certificate files, public DNS, PM2 certificate reload, Internet-facing TLS scan, and test accounts for every business role.")

    add_heading(doc, "3. Findings Remediation Matrix", 1)
    rows = [
        ("1", "Plain HTTP / sensitive transmission", "Code implemented; deployment verification required", "Production Node HTTPS server, TLS 1.2 minimum, fixed HTTPS origin, separate HTTP 308 redirect. Startup fails without valid certificate, key, and public origin."),
        ("2", "CSRF", "Implemented", "Session-bound CSRF token endpoint; state-changing API requests require X-CSRF-Token; non-whitelisted Origin is rejected."),
        ("3", "Parameter tampering", "Implemented", "Zod DTO validation, unknown-field rejection, ID and range checks; server recomputes authorization from authenticated user rather than trusting request fields."),
        ("4", "Weak administrator password", "Implemented", "Minimum 8 characters and at least 3 character classes; bcrypt hashing with 12 rounds; no fixed default password."),
        ("5", "Missing security headers", "Implemented", "Helmet CSP and related headers; HSTS in production; no-store for sensitive responses; anti-framing policy."),
        ("6", "Horizontal privilege escalation", "Implemented", "Unified read/manage/update checks for objectives, KRs, comments, notifications, import/export, analytics, and AI analysis. Unauthorized resources return 404."),
        ("7", "Vertical privilege escalation", "Implemented", "Administrator routes retain server-side super-administrator checks; client route guard redirects non-administrators away from management pages."),
        ("8", "Insufficient audit", "Implemented", "PostgreSQL audit_logs records actor, role, time, IP, user agent, action, resource, result, and request ID; restricted administrator query/export; 180-day retention design."),
    ]
    add_table(doc, ["No.", "Finding", "Status", "Remediation evidence"], rows, [520, 1750, 2200, 4890])

    add_heading(doc, "4. Key Technical Controls", 1)
    add_heading(doc, "4.1 Authorization Model", 2)
    add_para(doc, "The authorization model is enforced on the API layer and does not rely on UI visibility. It uses the authenticated user's role plus their relationship to the requested resource.")
    add_table(doc, ["Identity", "Permitted scope"], [
        ("Super administrator", "Manage all resources and system configuration."),
        ("Objective creator", "Manage objectives they created and their KRs."),
        ("KR assignee", "View the associated objective; update own KR progress and self-assessment only."),
        ("KR collaborator", "View the associated objective and KR; add comments."),
        ("Department member / leadership", "Read and comment only within the existing visibility scope; cannot manage resources created by others."),
        ("Non-administrator", "Uses DingTalk authentication only; local username/password login is rejected."),
    ], [2400, 6960])
    add_para(doc, "For resource IDs outside a user's scope, the API returns 404 to reduce resource enumeration. Access to administrator-only modules returns 403 after server-side role validation.", color=DARK_BLUE)

    add_heading(doc, "4.2 Authentication and Session Security", 2)
    add_bullet(doc, "Only a local super-administrator account may use POST /api/auth/login. Other users authenticate through DingTalk.")
    add_bullet(doc, "Existing passwords are bcrypt hashes in the database. Legacy AES-based administrator bootstrap handling and hard-coded AES key logic were removed.")
    add_bullet(doc, "A production deployment requires SESSION_SECRET. There is no fixed fallback secret. The session ID is regenerated after successful local login.")
    add_bullet(doc, "Login protection applies both IP- and account-dimension rate limiting, with temporary lockout after consecutive failed attempts.")
    add_bullet(doc, "Production session cookies use Secure, HttpOnly, and SameSite=Lax. Direct TLS mode does not set proxy trust by default.")

    add_heading(doc, "4.3 Input, Browser, and Response Protections", 2)
    add_bullet(doc, "Write APIs use Zod DTOs and field whitelists. Checks cover identifiers, dates, role/status values, 0-100 progress, scores, weights, arrays, and import row limits.")
    add_bullet(doc, "The client obtains the CSRF token through GET /api/auth/csrf-token and attaches it as X-CSRF-Token to state-changing requests.")
    add_bullet(doc, "Helmet configuration supplies CSP, frame-ancestors, X-Content-Type-Options, Referrer-Policy, and production HSTS. Sensitive pages are marked Cache-Control: no-store.")

    add_heading(doc, "4.4 Audit and Logging", 2)
    add_para(doc, "Security-relevant operations are written to audit_logs, including login success/failure, logout, password changes, user/role/department/cycle changes, objective and KR operations, progress and scoring, import/export, AI analysis, and bulk data clearing.")
    add_para(doc, "Audit records intentionally exclude passwords, password hashes, session values, CSRF tokens, cookies, full request bodies, certificate paths, and private-key information. Routine request logging is structured and limited to request ID, method, path, status, and duration.")

    add_heading(doc, "5. HTTPS Deployment Requirements", 1)
    add_para(doc, "The code supports direct Node.js TLS termination. The following items are mandatory for production deployment and are test-environment acceptance prerequisites:")
    add_table(doc, ["Configuration", "Required behavior"], [
        ("HTTPS_CERT_PATH", "Readable PEM full certificate chain."),
        ("HTTPS_KEY_PATH", "Readable PEM private key; file permission limited to the process user."),
        ("HTTPS_PORT", "HTTPS listener; default 5000."),
        ("HTTP_PORT", "HTTP redirect listener; default 5001."),
        ("PUBLIC_HTTPS_ORIGIN", "Fixed external HTTPS origin. Redirects must not use the Host header."),
        ("Failure behavior", "In production, missing, unreadable, invalid, or mismatched TLS material stops startup; no HTTP fallback."),
        ("Renewal", "After certificate renewal, run pm2 reload okr-platform; in-process hot certificate reload is not used."),
    ], [2500, 6860])

    add_heading(doc, "6. Test Regression Checklist", 1)
    add_heading(doc, "6.1 Authorization", 2)
    for item in [
        "Use accounts for super administrator, objective creator, KR assignee, collaborator, same-department reader, and cross-department reader.",
        "Tamper objectiveId, KR ID, comment ID, notification ID, and reorder request IDs. Confirm cross-scope write/delete is rejected and cross-scope resource lookup returns 404.",
        "Confirm a KR assignee cannot modify another person's KR, objective creator, department, collaborator scope, or ordering outside their managed objective.",
        "Confirm notification read status can only be changed by the receiving user; only comment author or super administrator can delete a comment.",
        "Pass an arbitrary department parameter to AI analysis/ranking. Confirm results are limited to the current user's visible objectives.",
    ]:
        add_number(doc, item)
    add_heading(doc, "6.2 Authentication and Browser Security", 2)
    for item in [
        "Attempt local password login for a non-administrator. Confirm a generic credential failure response.",
        "Attempt weak administrator password changes. Confirm policy rejection. Attempt repeated failures by account and IP; confirm temporary lockout.",
        "Verify successful login changes the session identifier. Verify cookies have Secure, HttpOnly, and SameSite=Lax in production HTTPS mode.",
        "Send state-changing requests without X-CSRF-Token, with an invalid token, and from a non-whitelisted Origin. Confirm rejection.",
        "Verify response headers include CSP, anti-framing, X-Content-Type-Options, Referrer-Policy, and HSTS in production.",
    ]:
        add_number(doc, item)
    add_heading(doc, "6.3 TLS and Operations", 2)
    for item in [
        "Start production mode with a valid PEM certificate/key pair; verify direct HTTPS login and API behavior.",
        "Verify startup fails for missing certificate, missing key, unreadable file, invalid PEM, and mismatched certificate/key.",
        "Verify TLS 1.0 and 1.1 are rejected. Verify HTTP returns a 308 redirect preserving path and query parameters.",
        "Verify Host-header injection does not alter the redirect target. Verify pm2 reload okr-platform loads renewed certificate material.",
        "Verify audit records are created for login, password change, user/role changes, objective/KR changes, import/export, AI analysis, and data clear operations; verify sensitive values are absent.",
    ]:
        add_number(doc, item)

    add_heading(doc, "7. Release Conditions and Open Items", 1)
    add_table(doc, ["Release condition", "Owner", "Required evidence"], [
        ("Complete multi-role security regression in test environment", "Test", "Executed checklist, API request/response evidence, and defect closure record."),
        ("Deploy TLS certificates and production variables", "Operations", "PM2 configuration, restricted certificate file permissions, HTTPS endpoint evidence."),
        ("Rotate exposed credentials and remove legacy secrets", "Operations / Administrator", "New secrets recorded in secret management; legacy AES bootstrap value removed; sessions invalidated after rotation."),
        ("Verify audit retention job and monitoring", "Backend / Operations", "180-day cleanup job record and monitoring alerts for TLS startup failure, login failures, authorization denials, exceptions, and audit-write failures."),
        ("Review third-party dependency findings", "Backend", "Dependency upgrade/mitigation plan; document residual risk for packages without an upstream fix."),
    ], [3300, 1600, 4460])
    add_para(doc, "The application must not be released to production until the above release conditions are completed and the security test team issues a passed retest result.", bold_prefix="The application must not be released to production", color=RED)

    add_heading(doc, "8. Appendix: Test-Facing API Changes", 1)
    add_table(doc, ["Interface / behavior", "Test expectation"], [
        ("GET /api/auth/csrf-token", "Returns the current session's CSRF token."),
        ("POST /api/auth/login", "Local login is accepted only for a local super administrator; invalid credentials produce a generic failure response."),
        ("GET /api/auth/me", "Returns authProvider and excludes password-related fields."),
        ("State-changing APIs", "Require X-CSRF-Token and accepted Origin; reject unknown payload fields."),
        ("Resource authorization", "Unauthorized resource IDs return 404; administrator-only module access returns 403."),
    ], [3000, 6360])

    doc.core_properties.title = "OKR Platform Security Remediation Report"
    doc.core_properties.subject = "Security regression testing evidence"
    doc.core_properties.author = "OKR Platform Engineering"
    doc.core_properties.comments = "Generated for security retest."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
